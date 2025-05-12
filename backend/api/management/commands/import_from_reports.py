# api/management/commands/import_from_reports.py
import re, zipfile
from pathlib import Path
from docx import Document
from django.core.management.base import BaseCommand
from api.models import Student, Curator, Skill, StudentSkill

RX_CUR  = re.compile(r'(?i)научн\w*\s+руководител[^\n]*?[–:]\s*(.+)')
RX_TECH = re.compile(r'(?i)технолог')
BULLETS = r'[•▪·\-—–]'

def split_names(raw):  # Ivanov I.I., Petrov P.P.
    raw = raw.replace(' и ', ',')
    return [n.strip() for n in re.split(r'[;,]', raw) if n.strip()]

def extract_students(texts):
    for i, t in enumerate(texts):
        if re.search(r'(?i)выполн\w*$', t.strip()):
            out = []
            for ln in texts[i+1:]:
                ln = ln.strip()
                if not ln or re.search(r'(?i)научн', ln):
                    break
                if not re.search(r'курс|групп', ln):
                    out.append(ln)
            return out
    return []

def extract_skills(doc):
    skills = []
    idx = next((i for i, p in enumerate(doc.paragraphs)
                if RX_TECH.search(p.text)), None)
    if idx is None:
        return skills
    for p in doc.paragraphs[idx+1:]:
        if p.style.name.lower().startswith('heading'):
            break
        txt = p.text.strip()
        if not txt:
            continue
        if re.match(BULLETS, txt):
            txt = re.sub(BULLETS, '', txt, 1).strip()
        skills += [s.strip() for s in re.split(r'[;,]', txt) if s.strip()]
    if skills:
        return skills
    for tbl in doc.tables:
        if any(RX_TECH.search(c.text) for row in tbl.rows for c in row.cells):
            ridx = next(i for i, row in enumerate(tbl.rows)
                        if any(RX_TECH.search(c.text) for c in row.cells))
            if ridx + 1 < len(tbl.rows):
                for cell in tbl.rows[ridx+1].cells:
                    skills += [s.strip() for s in re.split(r'[;,]', cell.text) if s.strip()]
            break
    return skills

class Command(BaseCommand):
    help = "Импорт студентов и их навыков из отчётов .docx"

    def add_arguments(self, parser):
        parser.add_argument("path", help="Файл .docx или папка")

    def handle(self, *args, path, **opts):
        paths = [Path(path)] if Path(path).is_file() else Path(path).glob("*.doc*")
        paths = list(paths)
        self.stdout.write(f"Файлов найдено: {len(paths)}")

        new_students = new_skills = skipped = 0

        for file in paths:
            try:
                doc = Document(file)
            except zipfile.BadZipFile:
                skipped += 1
                self.stdout.write(self.style.WARNING(
                    f"⨯ Пропущен {file.name} (не .docx или повреждён)"))
                continue

            # абзацы + таблицы
            texts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        texts.append(cell.text)

            students = extract_students(texts)

            m_cur = next((RX_CUR.search(t) for t in texts if RX_CUR.search(t)), None)
            curator = None
            if m_cur:
                curator_name = m_cur.group(1).strip()
                curator, _ = Curator.objects.get_or_create(
                    name=curator_name, defaults={'email': ''}
                )

            skills = extract_skills(doc)

            for stud in students:
                stu, created = Student.objects.get_or_create(name=stud, defaults={'email': ''})
                if created:
                    new_students += 1
                for sk in skills:
                    skill, created = Skill.objects.get_or_create(name=sk)
                    if created:
                        new_skills += 1
                    StudentSkill.objects.update_or_create(
                        student=stu, skill=skill, defaults={'level': 1.0}
                    )

        self.stdout.write(self.style.SUCCESS(
            f"\nИмпорт завершён → новых студентов: {new_students}, "
            f"новых навыков: {new_skills}, пропущено файлов: {skipped}"
        ))
